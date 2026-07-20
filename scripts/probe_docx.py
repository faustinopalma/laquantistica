import sys, traceback
from pathlib import Path

root = Path(__file__).resolve().parent.parent

files = [
    "1. Esperimento di Stern-Gerlach/Esperimento di Stern-Gerlach.docx",
    "2. Esperimenti di Stern-Gerlach in cascata/Esperimenti di Stern-Gerlach in cascata.docx",
    "3. Esperimenti con gli Elettroni/ESPERIMENTI CON GLI ELETTRONI.docx",
    "3. Esperimenti con gli Elettroni/esperimenti_con_gli_elettroni_convertito.docx",
    "4. Diffrazione degli Elettroni/DIFFRAZIONE DEGLI ELETTRONI.docx",
    "5. Esperimento di Rutherford/ESPERIMENTO DI RUTHERFORD 2.docx",
    "5. Esperimento di Rutherford/ESPERIMENTO DI RUTHERFORD.docx",
    "6. Ulteriori sviluppi della Teoria/Ulteriori sviluppi della Teoria.docx",
    "7. Esperimento di Franck-Hertz/ESPERIMENTO DI FRANCK-HERTZ.docx",
    "8. Effetto Fotoelettrico/EFFETTO FOTOELETTRICO.docx",
    "9. Spettri atomici di emissione/SPETTRI ATOMICI DI EMISSIONE.docx",
    "Introduzione.docx",
]

from docx import Document

for f in files:
    p = root / f
    print("=" * 80)
    print(f)
    if not p.exists():
        print("  MISSING")
        continue
    try:
        doc = Document(str(p))
        n_par = len(doc.paragraphs)
        text = "\n".join(par.text for par in doc.paragraphs)
        n_char = len(text.strip())
        print(f"  OK paragraphs={n_par} chars={n_char}")
        print("  first 200:", repr(text.strip()[:200]))
    except Exception as e:
        print("  ERROR:", type(e).__name__, e)
